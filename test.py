from docx import Document

doc = Document()
doc.add_heading("Employment Application Form", level=1)

doc.add_paragraph("Please complete all required fields.\n")

fields = [
    "Full Name",
    "Date of Birth",
    "Email Address",
    "Phone Number",
    "Home Address",
    "Position Applied For",
    "Date of Application",
    "Signature"
]

for field in fields:
    doc.add_paragraph(f"{field}: ___________________________")

doc.save("form2.docx")
