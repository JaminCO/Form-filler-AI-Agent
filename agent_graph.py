import os
from typing import Dict, List, Any, Optional, TypedDict, Annotated
from pathlib import Path
import json
import re

# LangGraph and LangChain imports
from langgraph.graph import StateGraph, END
from langgraph.prebuilt import ToolExecutor
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain.prompts import ChatPromptTemplate

# Document processing
import docx
from docxtpl import DocxTemplate
import fitz  # PyMuPDF
import pandas as pd
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class AgentState(TypedDict):
    """State that gets passed between nodes in the LangGraph"""
    form_path: str
    knowledge_base_paths: List[str]
    extracted_fields: List[Dict[str, Any]]
    knowledge_data: List[Document]
    filled_fields: Dict[str, Any]
    missing_fields: List[str]
    user_questions: List[str]
    user_answers: Dict[str, str]
    output_path: Optional[str]
    current_step: str
    error_message: Optional[str]

class FormFillingAgent:
    def __init__(self):
        """Initialize the Form Filling Agent with LLM and embeddings"""
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-pro",
            google_api_key=os.getenv("GOOGLE_API_KEY"),
            temperature=0.1
        )
        
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=os.getenv("GOOGLE_API_KEY")
        )
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
        
        # Initialize the LangGraph workflow
        self.workflow = self._build_workflow()
        
    def _build_workflow(self) -> StateGraph:
        """Build the LangGraph workflow with nodes and edges"""
        workflow = StateGraph(AgentState)
        
        # Add nodes
        workflow.add_node("parse_form", self.parse_form_node)
        workflow.add_node("process_knowledge_base", self.process_knowledge_base_node)
        workflow.add_node("extract_and_match", self.extract_and_match_node)
        workflow.add_node("llm_reasoning", self.llm_reasoning_node)
        workflow.add_node("generate_questions", self.generate_questions_node)
        workflow.add_node("fill_form", self.fill_form_node)
        workflow.add_node("handle_error", self.handle_error_node)
        
        # Define the workflow edges
        workflow.set_entry_point("parse_form")
        
        workflow.add_edge("parse_form", "process_knowledge_base")
        workflow.add_edge("process_knowledge_base", "extract_and_match")
        workflow.add_edge("extract_and_match", "llm_reasoning")
        
        # Conditional edge: if missing fields, ask questions; otherwise fill form
        workflow.add_conditional_edges(
            "llm_reasoning",
            self.should_ask_questions,
            {
                "ask_questions": "generate_questions",
                "fill_form": "fill_form",
                "error": "handle_error"
            }
        )
        
        workflow.add_edge("generate_questions", END)
        workflow.add_edge("fill_form", END)
        workflow.add_edge("handle_error", END)
        
        return workflow.compile()

    def parse_form_node(self, state: AgentState) -> AgentState:
        """Parse the uploaded form and extract fields"""
        try:
            form_path = Path(state["form_path"])
            extracted_fields = []
            
            if form_path.suffix.lower() == '.pdf':
                extracted_fields = self._parse_pdf_form(form_path)
            elif form_path.suffix.lower() in ['.docx', '.doc']:
                extracted_fields = self._parse_docx_form(form_path)
            else:
                state["error_message"] = f"Unsupported form format: {form_path.suffix}"
                return state
            
            state["extracted_fields"] = extracted_fields
            state["current_step"] = "form_parsed"
            
        except Exception as e:
            state["error_message"] = f"Error parsing form: {str(e)}"
            
        return state

    def _parse_pdf_form(self, pdf_path: Path) -> List[Dict[str, Any]]:
        """Extract form fields from PDF"""
        fields = []
        doc = fitz.open(pdf_path)
        
        # Check if PDF has form fields
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            # Get form fields if they exist
            if page.widgets():
                for widget in page.widgets():
                    fields.append({
                        "name": widget.field_name or f"field_{len(fields)}",
                        "type": widget.field_type_string,
                        "value": widget.field_value or "",
                        "required": widget.field_flags & 2 != 0,
                        "page": page_num
                    })
            else:
                # If no form fields, extract text and identify potential fields
                text = page.get_text()
                potential_fields = self._identify_fields_from_text(text)
                fields.extend(potential_fields)
        
        doc.close()
        return fields

    def _parse_docx_form(self, docx_path: Path) -> List[Dict[str, Any]]:
        """Extract form fields from DOCX"""
        doc = docx.Document(docx_path)
        fields = []
        
        # Extract text and look for form patterns
        full_text = ""
        for paragraph in doc.paragraphs:
            full_text += paragraph.text + "\n"
        
        # Look for underlines, brackets, or other field indicators
        fields = self._identify_fields_from_text(full_text)
        
        return fields

    def _identify_fields_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Use LLM to identify form fields from text"""
        prompt = ChatPromptTemplate.from_template(
            """
            Analyze the following form text and identify all the fields that need to be filled out.
            Extract field names, determine if they're required, and classify the type of information needed.
            
            Form text:
            {text}
            
            Return a JSON list of fields with this structure:
            [
                {
                    "name": "field_name",
                    "type": "text|date|number|email|address|phone",
                    "description": "what information is being asked for",
                    "required": true|false,
                    "context": "surrounding text for context"
                }
            ]
            
            Focus on fields like: name, address, phone, email, date of birth, SSN, employment info, etc.
            """
        )
        
        try:
            response = self.llm.invoke(prompt.format(text=text[:3000]))  # Limit text length
            
            # Try to parse JSON from response
            content = response.content
            if "```json" in content:
                json_str = content.split("```json")[1].split("```")[0]
            else:
                json_str = content
                
            fields = json.loads(json_str)
            return fields
        except Exception as e:
            print(f"Error identifying fields: {e}")
            return []

    def process_knowledge_base_node(self, state: AgentState) -> AgentState:
        """Process and vectorize the knowledge base documents"""
        try:
            documents = []
            
            for kb_path in state["knowledge_base_paths"]:
                kb_path = Path(kb_path)
                
                if kb_path.suffix.lower() == '.pdf':
                    docs = self._extract_pdf_text(kb_path)
                elif kb_path.suffix.lower() in ['.docx', '.doc']:
                    docs = self._extract_docx_text(kb_path)
                elif kb_path.suffix.lower() == '.txt':
                    with open(kb_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    docs = [Document(page_content=text, metadata={"source": str(kb_path)})]
                
                documents.extend(docs)
            
            # Split documents into chunks
            split_docs = self.text_splitter.split_documents(documents)
            
            state["knowledge_data"] = split_docs
            state["current_step"] = "knowledge_processed"
            
        except Exception as e:
            state["error_message"] = f"Error processing knowledge base: {str(e)}"
            
        return state

    def _extract_pdf_text(self, pdf_path: Path) -> List[Document]:
        """Extract text from PDF"""
        doc = fitz.open(pdf_path)
        documents = []
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                documents.append(Document(
                    page_content=text,
                    metadata={"source": str(pdf_path), "page": page_num}
                ))
        
        doc.close()
        return documents

    def _extract_docx_text(self, docx_path: Path) -> List[Document]:
        """Extract text from DOCX"""
        doc = docx.Document(docx_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return [Document(
            page_content=text,
            metadata={"source": str(docx_path)}
        )]

    def extract_and_match_node(self, state: AgentState) -> AgentState:
        """Match form fields with knowledge base data using vector search"""
        try:
            if not state["knowledge_data"]:
                state["error_message"] = "No knowledge data available"
                return state
            
            # Create vector store
            vectorstore = FAISS.from_documents(
                state["knowledge_data"], 
                self.embeddings
            )
            
            filled_fields = {}
            
            for field in state["extracted_fields"]:
                field_name = field["name"]
                field_description = field.get("description", field_name)
                
                # Search for relevant information
                query = f"{field_name} {field_description}"
                relevant_docs = vectorstore.similarity_search(query, k=3)
                
                if relevant_docs:
                    # Use LLM to extract specific information
                    context = "\n".join([doc.page_content for doc in relevant_docs])
                    extracted_value = self._extract_field_value(field, context)
                    
                    if extracted_value:
                        filled_fields[field_name] = extracted_value
            
            state["filled_fields"] = filled_fields
            state["current_step"] = "fields_matched"
            
        except Exception as e:
            state["error_message"] = f"Error matching fields: {str(e)}"
            
        return state

    def _extract_field_value(self, field: Dict[str, Any], context: str) -> Optional[str]:
        """Use LLM to extract specific field value from context"""
        prompt = ChatPromptTemplate.from_template(
            """
            Extract the value for the field "{field_name}" from the following context.
            
            Field description: {field_description}
            Field type: {field_type}
            
            Context:
            {context}
            
            Rules:
            - Return ONLY the extracted value, no explanation
            - If the information is not found, return "NOT_FOUND"
            - Format the value appropriately (e.g., dates as MM/DD/YYYY)
            - For addresses, return full address on one line
            
            Value:
            """
        )
        
        try:
            response = self.llm.invoke(prompt.format(
                field_name=field["name"],
                field_description=field.get("description", ""),
                field_type=field.get("type", "text"),
                context=context[:2000]  # Limit context length
            ))
            
            value = response.content.strip()
            return None if value == "NOT_FOUND" else value
            
        except Exception as e:
            print(f"Error extracting field value: {e}")
            return None

    def llm_reasoning_node(self, state: AgentState) -> AgentState:
        """Use LLM reasoning to fill missing fields or infer information"""
        try:
            missing_fields = []
            
            for field in state["extracted_fields"]:
                field_name = field["name"]
                
                if field_name not in state["filled_fields"]:
                    # Try to infer the value using LLM reasoning
                    inferred_value = self._infer_field_value(field, state["filled_fields"])
                    
                    if inferred_value:
                        state["filled_fields"][field_name] = inferred_value
                    else:
                        missing_fields.append(field_name)
            
            state["missing_fields"] = missing_fields
            state["current_step"] = "llm_reasoning_complete"
            
        except Exception as e:
            state["error_message"] = f"Error in LLM reasoning: {str(e)}"
            
        return state

    def _infer_field_value(self, field: Dict[str, Any], filled_fields: Dict[str, Any]) -> Optional[str]:
        """Try to infer field value based on other filled fields"""
        if not filled_fields:
            return None
        
        prompt = ChatPromptTemplate.from_template(
            """
            Based on the following information, try to infer or generate a reasonable value for the field "{field_name}".
            
            Field description: {field_description}
            Field type: {field_type}
            
            Available information:
            {available_info}
            
            Rules:
            - Only provide a value if you can reasonably infer it from the available information
            - For dates, use MM/DD/YYYY format
            - For phone numbers, use (XXX) XXX-XXXX format
            - If you cannot infer the value, return "CANNOT_INFER"
            
            Inferred value:
            """
        )
        
        try:
            available_info = "\n".join([f"{k}: {v}" for k, v in filled_fields.items()])
            
            response = self.llm.invoke(prompt.format(
                field_name=field["name"],
                field_description=field.get("description", ""),
                field_type=field.get("type", "text"),
                available_info=available_info
            ))
            
            value = response.content.strip()
            return None if value == "CANNOT_INFER" else value
            
        except Exception as e:
            print(f"Error inferring field value: {e}")
            return None

    def generate_questions_node(self, state: AgentState) -> AgentState:
        """Generate clarifying questions for missing fields"""
        try:
            questions = []
            
            for field_name in state["missing_fields"]:
                field = next((f for f in state["extracted_fields"] if f["name"] == field_name), None)
                
                if field:
                    question = self._generate_field_question(field)
                    questions.append(question)
            
            state["user_questions"] = questions
            state["current_step"] = "questions_generated"
            
        except Exception as e:
            state["error_message"] = f"Error generating questions: {str(e)}"
            
        return state

    def _generate_field_question(self, field: Dict[str, Any]) -> str:
        """Generate a user-friendly question for a missing field"""
        field_name = field["name"]
        field_type = field.get("type", "text")
        description = field.get("description", "")
        
        if field_type == "email":
            return f"What is your email address for {field_name}?"
        elif field_type == "phone":
            return f"What is your phone number for {field_name}?"
        elif field_type == "date":
            return f"What is the date for {field_name}? (Please use MM/DD/YYYY format)"
        elif field_type == "address":
            return f"What is your full address for {field_name}?"
        else:
            return f"Please provide the information for {field_name}" + (f" ({description})" if description else "")

    def fill_form_node(self, state: AgentState) -> AgentState:
        """Fill the form with collected data and generate output"""
        try:
            form_path = Path(state["form_path"])
            output_path = form_path.parent / f"completed_{form_path.name}"
            
            # Merge user answers with filled fields
            all_data = {**state["filled_fields"], **state.get("user_answers", {})}
            
            if form_path.suffix.lower() == '.pdf':
                self._fill_pdf_form(form_path, output_path, all_data)
            elif form_path.suffix.lower() in ['.docx', '.doc']:
                self._fill_docx_form(form_path, output_path, all_data)
            
            state["output_path"] = str(output_path)
            state["current_step"] = "form_completed"
            
        except Exception as e:
            state["error_message"] = f"Error filling form: {str(e)}"
            
        return state

    def _fill_pdf_form(self, input_path: Path, output_path: Path, data: Dict[str, Any]):
        """Fill PDF form fields"""
        doc = fitz.open(input_path)
        
        for page_num in range(doc.page_count):
            page = doc[page_num]
            
            for widget in page.widgets():
                field_name = widget.field_name
                if field_name in data:
                    widget.field_value = str(data[field_name])
                    widget.update()
        
        doc.save(output_path)
        doc.close()

    def _fill_docx_form(self, input_path: Path, output_path: Path, data: Dict[str, Any]):
        """Fill DOCX template with data"""
        try:
            # Try using docxtpl for template replacement
            template = DocxTemplate(input_path)
            template.render(data)
            template.save(output_path)
        except Exception as e:
            # Fallback: simple text replacement
            doc = docx.Document(input_path)
            
            for paragraph in doc.paragraphs:
                for field_name, value in data.items():
                    if field_name.lower() in paragraph.text.lower():
                        # Simple replacement - this could be enhanced
                        paragraph.text = paragraph.text.replace(f"[{field_name}]", str(value))
                        paragraph.text = paragraph.text.replace(f"{{{field_name}}}", str(value))
            
            doc.save(output_path)

    def should_ask_questions(self, state: AgentState) -> str:
        """Determine next step based on missing fields"""
        if state.get("error_message"):
            return "error"
        elif state.get("missing_fields") and not state.get("user_answers"):
            return "ask_questions"
        else:
            return "fill_form"

    def handle_error_node(self, state: AgentState) -> AgentState:
        """Handle errors in the workflow"""
        state["current_step"] = "error"
        return state

    def run_workflow(self, form_path: str, knowledge_base_paths: List[str], user_answers: Dict[str, str] = None) -> Dict[str, Any]:
        """Run the complete form filling workflow"""
        
        initial_state = AgentState(
            form_path=form_path,
            knowledge_base_paths=knowledge_base_paths,
            extracted_fields=[],
            knowledge_data=[],
            filled_fields={},
            missing_fields=[],
            user_questions=[],
            user_answers=user_answers or {},
            output_path=None,
            current_step="starting",
            error_message=None
        )
        
        # Run the workflow
        result = self.workflow.invoke(initial_state)
        
        return {
            "success": not result.get("error_message"),
            "error": result.get("error_message"),
            "output_path": result.get("output_path"),
            "questions": result.get("user_questions", []),
            "filled_fields": result.get("filled_fields", {}),
            "missing_fields": result.get("missing_fields", []),
            "step": result.get("current_step")
        }

# Example usage
if __name__ == "__main__":
    # Initialize the agent
    agent = FormFillingAgent()
    
    # Example: Run the workflow
    result = agent.run_workflow(
        form_path="form.docx",
        knowledge_base_paths=["knowledge_base/data.pdf", "knowledge_base/data2.pdf", "knowledge_base/data3.pdf"],
    )
    
    print("Workflow Result:")
    print(f"Success: {result['success']}")
    
    if result['questions']:
        print("\nQuestions for user:")
        for q in result['questions']:
            print(f"- {q}")
    
    if result['output_path']:
        print(f"\nCompleted form saved to: {result['output_path']}")