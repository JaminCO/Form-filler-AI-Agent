from langchain_community.document_loaders import PyPDFLoader
from langchain.chat_models import init_chat_model
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_core.prompts import PromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
# from langchain_pinecone import PineconeVectorStore
from langgraph.graph import StateGraph, START, END

import os
import time
from typing import Dict, List, Any, Optional, TypedDict, Annotated
from pathlib import Path
import json
import re

from pinecone import Pinecone
from docx import Document
import docx
from docxtpl import DocxTemplate
import fitz  # PyMuPDF
import pandas as pd
import PyPDF2
from dotenv import load_dotenv
from langchain_core.vectorstores import InMemoryVectorStore
import faiss

load_dotenv()


if not os.environ.get("GOOGLE_API_KEY"):
  print("Please set the GOOGLE_API_KEY environment variable.")


# pc = Pinecone(api_key=os.environ.get("PINECONE_API_KEY"))
# index = pc.Index(index_name)
# vector_store = PineconeVectorStore(embedding=embeddings, index=index)


embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
# vector_store = InMemoryVectorStore(embeddings)
llm = init_chat_model("gemini-2.5-flash", model_provider="google_genai")

text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200
)

class State(TypedDict):
    """State that gets passed between nodes in the LangGraph"""
    form_path: str
    knowledge_base_paths: List[str]
    extracted_fields: List[Dict[str, Any]]
    knowledge_data: List[Any]
    filled_fields: Dict[str, Any]
    missing_fields: List[str]
    user_questions: List[str]
    user_answers: Dict[str, str]
    output_path: Optional[str]
    current_step: str
    error_message: Optional[str]
    possible_answers: Optional[List[Dict]]


def parse_file(state: State) -> State:
    try:
        print("Parsing form...\n")
        form_path = state.get("form_path")
        extracted_fields = []

        if form_path.endswith(".docx"):
            extracted_fields = extract_form_data_docx(form_path)
        elif form_path.endswith(".pdf"):
            raise NotImplementedError("PDF parsing not implemented in this example.")
        else:
            state["error_message"] = f"Unsupported form format: {form_path.suffix}"
            return state

        state["extracted_fields"] = extracted_fields
        state["current_step"] = "form_parsed"
    
    except Exception as e:
        print(e)
        state["error_message"] = f"Error parsing form: {str(e)}"
    
    return state

def extract_form_data_docx(docx_path):
    doc = Document(docx_path)  
    fields = []

    full_text = ""
    for paragraph in doc.paragraphs:
        full_text += paragraph.text + "\n"

    example = '''
    [
        {
            "name": "field_name",
            "type": "text|date|number|email|address|phone",
            "description": "what information is being asked for",
            "required": true|false,
            "context": "surrounding text for context"
        }
    ]
'''

    custom_rag_prompt = PromptTemplate.from_template(
            """
            Analyze the following form text and identify all the fields that need to be filled out.
            Extract field names, determine if they're required, and classify the type of information needed.
            
            Form text:
            {text}
            
            Return a JSON list of fields with this structure, return no other text just this JSON Format all in english:
            {example}
            
            Focus on fields like: name, address, phone, email, date of birth, SSN, employment info, etc.
            """
        )
    try:
        messages = custom_rag_prompt.invoke({"text": full_text, "example": example})
        response = llm.invoke(messages)

        # Try to parse JSON from response
        content = response.content
        if "```json" in content:
            json_str = content.split("```json")[1].split("```")[0]
        else:
            json_str = content
            
        fields = json.loads(json_str)
        return fields
    except Exception as e:
        print(f"\nError identifying fields: {e}")
        return []

def ai_form_processing_pdf(pdf_path):
    print(f"Processing PDF form: {pdf_path}\n")
    loader = PyPDFLoader(pdf_path)
    documents = loader.load()

    # Combine all page text
    full_text = "\n".join([page.page_content for page in documents])
    total_length = len(full_text)

    # Split into 4 parts (or fewer if text is short)
    chunk_size = total_length // 4 if total_length > 0 else 0
    chunks = [full_text[i:i + chunk_size] for i in range(0, total_length, chunk_size)]
    if len(chunks) > 4:
        chunks = chunks[:4]  # ensure max 4 parts

    fields = []

    example = '''
    [
        {
            "name": "field_name",
            "value": "extracted_value",
            "type": "text|date|number|email|address|phone",
            "description": "what information is being asked for",
            "required": true|false,
            "context": "surrounding text for context"
        }
    ]
    '''

    custom_rag_prompt = PromptTemplate.from_template(
        """
        Analyze the following document text and identify all the field values.
        Extract field names and values, determine if they're required, and classify the type of information needed.

        Document text:
        {text}
        
        Return a JSON list of fields with this structure:
        {example}
        
        Rules:
        - Return ONLY the JSON LIST, no explanation
        - Skip fields if information is missing (no empty values)
        - Format dates as MM/DD/YYYY
        - Addresses should be full and on one line
        """
    )

    try:
        for chunk in chunks:
            if not chunk.strip():
                continue
            messages = custom_rag_prompt.invoke({"text": chunk, "example": example})
            response = llm.invoke(messages)

            content = response.content
            print(f"Chunk response: {content}\n")
            if "```json" in content:
                json_str = content.split("```json")[1].split("```")[0]
            else:
                json_str = content

            try:
                partial_fields = json.loads(json_str)
                if isinstance(partial_fields, list):
                    fields.extend(partial_fields)
            except Exception as e:
                print(f"Error parsing JSON for chunk: {e}")
                continue

        return fields

    except Exception as e:
        print(f"\nError identifying fields: {e}")
        return []

def process_knowledge_base(state: State):
    try:
        documents = []

        forms_path = state["knowledge_base_paths"]
        for f in forms_path:
            if f.endswith(".docx"):
                extracted_doc = extract_kb_docx(f)
            elif f.endswith(".pdf"):
                extracted_doc = extract_kb_pdf(f)
            # elif kb_path.suffix.lower() == '.txt':
            #     with open(kb_path, 'r', encoding='utf-8') as f:
            #         text = f.read()
            #     docs = [Document(page_content=text, metadata={"source": str(kb_path)})]
            else:
                raise ValueError("Unsupported file format. Only .docx files are supported in this example.")   
            documents.extend(extracted_doc)
            new_res = ai_form_processing_pdf(f)
        # Split documents into chunks
        split_docs = text_splitter.split_documents(documents)

        state["knowledge_data"] = split_docs
        state["possible_answers"] = new_res
        print(f"\nProcessed {len(split_docs)} knowledge base documents.")
        print("\n")
        state["current_step"] = "knowledge_processed"

    except Exception as e:#
        print(e)
        state["error_message"] = f"Error processing knowledge base: {str(e)}"
            
    return state


def extract_kb_pdf(pdf_path):
    loader = PyPDFLoader(pdf_path)
    documents = loader.load()
    return documents

def extract_kb_docx(docx_path):
    
    doc = Document(docx_path)  
    
    # full_text = ""
    # for paragraph in doc.paragraphs:
    #     full_text += paragraph.text + "\n"
    
    return doc


    # [Document(page_content=full_text, metadata={"source": str(docx_path)})]

def get_possible_answers_for_field(state: State, field_name: str, field_description: str = "") -> List[Dict]:
    """
    Searches the state's possible_answers list for relevant answers
    based on the field name and optional description.
    """
    if not state.get("possible_answers"):
        print("No possible answers in state")
        return []
    
    results = []
    for ans in state["possible_answers"]:
        ans_name = ans.get("name", "").lower()
        ans_desc = ans.get("description", "").lower()
        
        if field_name.lower() in ans_name or field_name.lower() in ans_desc:
            results.append(ans)
        elif field_description and field_description.lower() in ans_desc:
            results.append(ans)

    return results

def extract_and_match_node(state: State):
    """Match form fields with knowledge base data using vector search"""
    try:
        print("Matching fields with knowledge base...\n")
        if not state["knowledge_data"]:
            state["error_message"] = "No knowledge data available"
            print("\nNo knowledge data available\n")
            return state
        
        # Create vector store
        vectorstore = FAISS.from_documents(
            state["knowledge_data"], 
            embeddings
        )
        print(f"Vector store contains {vectorstore.index.ntotal} documents.\n")
        filled_fields = {}
        
        for field in state["extracted_fields"]:
            field_name = field["name"]
            field_desc = field.get("description", "")

            # Step 1: Check possible_answers
            candidates = get_possible_answers_for_field(state, field_name, field_desc)
            if candidates:
                filled_fields[field_name] = candidates[0].get("value", None)
                continue

            # Step 2: FAISS similarity search
            query = f"{field_name} {field_desc}"
            relevant_docs = vectorstore.similarity_search(query, k=3)
            if relevant_docs:
                context = "\n".join([doc.page_content for doc in relevant_docs])
                extracted_value = get_kb_data(field, context)
                if extracted_value:
                    filled_fields[field_name] = extracted_value
                    continue

            # Step 3: Fallback — leave for llm_reasoning_node
            
        state["filled_fields"].update(filled_fields)
        print("\nFilled fields after matching:", filled_fields)
        print("\n")
        state["current_step"] = "fields_matched"
        
    except Exception as e:
        print(e)
        state["error_message"] = f"Error matching fields: {str(e)}"
        
    return state
    
def get_kb_data(field, context):
    custom_rag_prompt = PromptTemplate.from_template(
            """
            Analyze the following document text and identify all the fields values.
            Extract the value for the field name"{field_name}" from the following context.
            
            Field description: {field_description}
            Field type: {field_type}
            
            Context:
            {context}
            
            Rules:
            - Return ONLY the extracted value, no explanation
            - If the information is not found, return "NOT_FOUND"
            - Format the value appropriately (e.g., dates as MM/DD/YYYY)
            - For addresses, return full address on one line
            
            Value:
            """
        )
    try:
        data = {"field_name":field["name"], "field_description":field.get("description", ""), "field_type":field.get("type", "text"), "context":context}
        print(f"Getting data for field: {data}")
        messages = custom_rag_prompt.invoke({"field_name":field["name"], "field_description":field.get("description", ""), "field_type":field.get("type", "text"), "context":context})
        response = llm.invoke(messages)
        time.sleep(1)  # to avoid rate limits
        value = response.content.strip()
        print(value)

        if value == "NOT_FOUND":
            return None
        elif "```json" in value:
            json_str = value.split("```json")[1].split("```")[0]
        else:
            json_str = value

        value = json.loads(json_str)
        return value
            
    except Exception as e:
        print(f"Error extracting field value: {e}")
        print(e)
        return None

def llm_reasoning_node(state: State):
    """Use LLM reasoning to fill missing fields or infer information"""
    try:
        print("Using LLM reasoning to fill missing fields...\n")
        missing_fields = []
        
        for field in state["extracted_fields"]:
            field_name = field["name"]
            
            if field_name not in state["filled_fields"]:
                # Try to infer the value using LLM reasoning
                inferred_value = infer_field_value(field, state["filled_fields"], state)
                
                if inferred_value:
                    state["filled_fields"][field_name] = inferred_value
                else:
                    missing_fields.append(field_name)
        
        state["missing_fields"] = missing_fields
        print("Filled fields after LLM reasoning:", state["filled_fields"])
        print("Missing fields after LLM reasoning:", missing_fields)
        print("\n")
        state["current_step"] = "llm_reasoning_complete"
        
    except Exception as e:
        print(e)
        state["error_message"] = f"Error in LLM reasoning: {str(e)}"
        
    return state

def infer_field_value(field: Dict[str, Any], filled_fields: Dict[str, Any], state) -> Optional[str]:
        """Try to infer field value based on other filled fields"""
        print(f"Inferring value for field: {field['name']}")
        if not filled_fields:
            return None
        
        prompt = PromptTemplate.from_template(
            """
            Based on the following information, try to infer or generate a reasonable value for the field "{field_name}".
            
            Field description: {field_description}
            Field type: {field_type}
            
            Available information:
            {available_info}

            and

            Other Information that might match the field:
            {possible_answers}
            
            Rules:
            - Only provide a value if you can reasonably infer it from the available information, no explanations
            - Return ONLY the extracted value, no explanation
            - Format the value appropriately (e.g., dates as DD/MM/YYYY)
            - For phone numbers, use (XXX) XXX-XXXX format
            - If you cannot infer the value, return "CANNOT_INFER"
            
            Inferred value:
            """
        )
        
        try:
            available_info = "\n".join([f"{k}: {v}" for k, v in filled_fields.items()])
            
            response = llm.invoke(prompt.format(
                field_name=field["name"],
                field_description=field.get("description", ""),
                field_type=field.get("type", "text"),
                available_info=available_info,
                possible_answers=json.dumps(state.get("possible_answers", []))
            ))
            
            value = response.content.strip()
            print("\nInferred value:", end=" ")
            print(value)
            return None if value == "CANNOT_INFER" else value
            
        except Exception as e:
            print(f"Error inferring field value: {e}")
            return None

def generate_questions_node(state: State):
    """Generate clarifying questions for missing fields"""
    try:
        questions = []
        
        for field_name in state["missing_fields"]:
            field = next((f for f in state["extracted_fields"] if f["name"] == field_name), None)
            
            if field:
                question = generate_field_question(field)
                questions.append(question)
        
        state["user_questions"] = questions
        print("Generated questions for user:", questions)
        print("\n")
        state["current_step"] = "questions_generated"
        
    except Exception as e:
        state["error_message"] = f"Error generating questions: {str(e)}"
        
    return state

def generate_field_question(field: Dict[str, Any]) -> str:
    """Generate a user-friendly question for a missing field"""
    print(f"Generating question for field: {field['name']}")
    field_name = field["name"]
    field_type = field.get("type", "text")
    description = field.get("description", "")
    
    if field_type == "email":
        return f"What is your email address for {field_name}?"
    elif field_type == "phone":
        return f"What is your phone number for {field_name}?"
    elif field_type == "date":
        return f"What is the date for {field_name}? (Please use MM/DD/YYYY format)"
    elif field_type == "address":
        return f"What is your full address for {field_name}?"
    else:
        return f"Please provide the information for {field_name}" + (f" ({description})" if description else "")

def fill_form_node(state: State):
    """Fill the form with collected data and generate output"""
    print("Filling form with collected data...\n")
    try:
        form_path = Path(state["form_path"])
        output_path = form_path.parent / f"completed_{form_path.name}"
        
        # Merge user answers with filled fields
        all_data = {**state["filled_fields"], **state.get("user_answers", {})}
        
        if form_path.suffix.lower() == '.pdf':
            fill_pdf_form(form_path, output_path, all_data)
        elif form_path.suffix.lower() in ['.docx', '.doc']:
            fill_docx_form(form_path, output_path, all_data)
        
        state["output_path"] = str(output_path)
        state["current_step"] = "form_completed"
        
    except Exception as e:
        state["error_message"] = f"Error filling form: {str(e)}"
        
    return state

def fill_pdf_form(input_path: Path, output_path: Path, data: Dict[str, Any]):
    """Fill PDF form fields"""
    print( f"Filling PDF form: {input_path} with data: {data}\n")
    doc = fitz.open(input_path)
    
    for page_num in range(doc.page_count):
        page = doc[page_num]
        
        for widget in page.widgets():
            field_name = widget.field_name
            if field_name in data:
                widget.field_value = str(data[field_name])
                widget.update()
    
    doc.save(output_path)
    doc.close()

def fill_docx_form(input_path, output_path, data):
    """
    input_path: path to input DOCX file
    output_path: where to save filled DOCX
    data: dictionary of field_name -> value
          Example: {"company_name": "Aurify AI Ltd", "idno": "12345", "date": "15/09/2025"}
    """

    doc = Document(input_path)

    for p in doc.paragraphs:
        text = p.text
        if "___" in text:  # only process lines with blanks
            for key, value in data.items():
                # if context keyword exists in the same paragraph, replace the first blank
                if re.search(key, text, re.IGNORECASE):
                    text = re.sub(r"_{3,}", value, text, count=1)
            p.text = text  # update paragraph

    doc.save(output_path)
    print(f"Filled form saved to: {output_path}")

def should_ask_questions(state: State):
    """Determine next step based on missing fields"""
    if state.get("error_message"):
        return "error"
    elif state.get("missing_fields") and not state.get("user_answers"):
        return "ask_questions"
    else:
        return "fill_form"

def handle_error_node(state: State) -> State:
    """Handle errors in the workflow"""
    state["current_step"] = "error"
    return state

# Create graph
graph = StateGraph(State)

# Add node
graph.add_node("parse_form", parse_file)
graph.add_node("process_knowledge_base", process_knowledge_base)
graph.add_node("extract_and_match", extract_and_match_node)
graph.add_node("llm_reasoning", llm_reasoning_node)
graph.add_node("generate_questions", generate_questions_node)
graph.add_node("fill_form", fill_form_node)
graph.add_node("handle_error", handle_error_node)

# Run graph

initial_state = State(
    form_path="form.docx",
    knowledge_base_paths=["knowledge_base/data3.pdf"],
    extracted_fields=[],
    knowledge_data=[],
    filled_fields={},
    missing_fields=[],
    user_questions=[],
    user_answers={},
    output_path=None,
    current_step="starting",
    error_message=None,
    possible_answers=[]
)

graph.add_edge(START, "parse_form")
graph.add_edge("parse_form", "process_knowledge_base")
graph.add_edge("process_knowledge_base", "extract_and_match")
graph.add_edge("extract_and_match", "llm_reasoning")
graph.add_edge("llm_reasoning", "generate_questions")
# graph.add_conditional_edges(
#     "llm_reasoning",
#     should_ask_questions,
#     {
#         "ask_questions": "generate_questions",
#         "fill_form": "fill_form",
#         "error": "handle_error"
#     }
# )

graph.add_edge("generate_questions", "fill_form")
graph.add_edge("fill_form", END)

graph_final = graph.compile()

result = graph_final.invoke(initial_state)